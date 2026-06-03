from __future__ import annotations

import csv
import json
import mimetypes
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from bs4 import BeautifulSoup

from deep_research.schemas import SourceProvenance


class IngestionError(RuntimeError):
    """Raised when an input source cannot be read into research text."""


@dataclass(frozen=True)
class IngestedDocument:
    title: str
    url: str
    content: str
    provenance: SourceProvenance
    metadata: dict[str, object]


def ingest_local_paths(paths: Iterable[str | Path]) -> list[IngestedDocument]:
    documents: list[IngestedDocument] = []
    for raw_path in paths:
        path = Path(raw_path).expanduser().resolve()
        if path.is_dir():
            candidates = sorted(child for child in path.rglob("*") if child.is_file())
        else:
            candidates = [path]
        for candidate in candidates:
            if not candidate.exists():
                raise IngestionError(f"Local input does not exist: {candidate}")
            try:
                content = read_local_document(candidate)
            except IngestionError:
                if path.is_dir():
                    continue
                raise
            if not content.strip():
                continue
            documents.append(
                IngestedDocument(
                    title=candidate.stem,
                    url=candidate.as_uri(),
                    content=content,
                    provenance="local_file",
                    metadata={
                        "path": str(candidate),
                        "suffix": candidate.suffix.lower(),
                    },
                )
            )
    return documents


def ingest_mcp_manifest(path: str | Path | None) -> list[IngestedDocument]:
    if not path:
        return []
    manifest_path = Path(path).expanduser().resolve()
    payload = json.loads(manifest_path.read_text(encoding="utf-8"))
    sources = payload.get("sources", payload if isinstance(payload, list) else [])
    if not isinstance(sources, list):
        raise IngestionError("MCP manifest must be a list or an object with a `sources` list.")

    documents: list[IngestedDocument] = []
    for index, source in enumerate(sources, start=1):
        if not isinstance(source, dict):
            continue
        content = str(source.get("content") or "")
        source_path = source.get("path")
        if not content and source_path:
            content = read_local_document(Path(str(source_path)).expanduser().resolve())
        if not content.strip():
            continue
        title = str(source.get("title") or source.get("name") or f"MCP Source {index}")
        url = str(source.get("url") or source.get("uri") or f"mcp://source/{index}")
        documents.append(
            IngestedDocument(
                title=title,
                url=url,
                content=content,
                provenance="mcp",
                metadata={key: value for key, value in source.items() if key not in {"content"}},
            )
        )
    return documents


def read_local_document(path: Path) -> str:
    suffix = path.suffix.lower()
    mime_type = mimetypes.guess_type(str(path))[0] or ""
    prefix = path.read_bytes()[:4096]
    if _looks_like_pdf(prefix, suffix, mime_type):
        return _pdf_to_text(path)
    if _looks_like_docx(suffix, mime_type):
        return _docx_to_text(path)
    if _looks_like_spreadsheet(suffix, mime_type):
        return _xlsx_to_text(path)
    if _looks_like_csv(suffix, mime_type, prefix):
        return _csv_to_text(path)
    if _looks_like_html(prefix, suffix, mime_type):
        html = path.read_text(encoding="utf-8", errors="replace")
        return _html_to_text(html)
    if _looks_like_text(prefix, mime_type):
        return path.read_text(encoding="utf-8", errors="replace")
    raise IngestionError(f"Unsupported local input type: {path.suffix or 'unknown'}")


def _looks_like_pdf(prefix: bytes, suffix: str, mime_type: str) -> bool:
    return suffix == ".pdf" or mime_type == "application/pdf" or prefix.startswith(b"%PDF")


def _looks_like_docx(suffix: str, mime_type: str) -> bool:
    return suffix == ".docx" or mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


def _looks_like_spreadsheet(suffix: str, mime_type: str) -> bool:
    return suffix in {".xls", ".xlsx"} or "spreadsheet" in mime_type or "excel" in mime_type


def _looks_like_csv(suffix: str, mime_type: str, prefix: bytes) -> bool:
    if suffix == ".csv" or mime_type == "text/csv":
        return True
    sample = prefix.decode("utf-8", errors="ignore")
    lines = [line for line in sample.splitlines()[:5] if line.strip()]
    return len(lines) >= 2 and all("," in line or "\t" in line for line in lines)


def _looks_like_html(prefix: bytes, suffix: str, mime_type: str) -> bool:
    if suffix in {".html", ".htm"} or mime_type in {"text/html", "application/xhtml+xml"}:
        return True
    sample = prefix.decode("utf-8", errors="ignore").lower()
    return bool(re.search(r"<!doctype\s+html|<html[\s>]|<body[\s>]|<article[\s>]", sample))


def _looks_like_text(prefix: bytes, mime_type: str) -> bool:
    if mime_type.startswith("text/") or mime_type in {"application/json", "application/xml"}:
        return True
    if b"\x00" in prefix:
        return False
    sample = prefix.decode("utf-8", errors="ignore")
    return bool(sample.strip()) and len(sample) / max(len(prefix), 1) > 0.75


def _html_to_text(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")
    for node in soup.select("script, style, nav, header, footer, aside"):
        node.decompose()
    return re.sub(r"\s+", " ", soup.get_text(" ", strip=True)).strip()


def _csv_to_text(path: Path) -> str:
    lines: list[str] = []
    with path.open(newline="", encoding="utf-8", errors="replace") as handle:
        for row in csv.reader(handle):
            lines.append(" | ".join(cell.strip() for cell in row if cell.strip()))
    return "\n".join(line for line in lines if line)


def _xlsx_to_text(path: Path) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError as exc:
        raise IngestionError("openpyxl is required to read spreadsheet inputs.") from exc
    workbook = load_workbook(path, read_only=True, data_only=True)
    lines: list[str] = []
    for sheet in workbook.worksheets:
        lines.append(f"# Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(value).strip() for value in row if value is not None and str(value).strip()]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def _pdf_to_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise IngestionError("pypdf is required to read PDF inputs.") from exc
    reader = PdfReader(str(path))
    return "\n\n".join((page.extract_text() or "").strip() for page in reader.pages).strip()


def _docx_to_text(path: Path) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise IngestionError("python-docx is required to read DOCX inputs.") from exc
    document = Document(path)
    return "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
